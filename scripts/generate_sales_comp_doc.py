from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def usd(value: float) -> str:
    return f"${value:,.2f}"


def pct(value: float) -> str:
    return f"{value * 100:.2f}%"


def build_document(output_path: str) -> None:
    doc = Document()

    title = doc.add_heading("Sales Compensation Foundations", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("A beginner-friendly guide with correct calculations and practical examples")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    doc.add_heading("1) What Is Sales Compensation?", level=1)
    doc.add_paragraph(
        "Sales compensation is how a company pays sales employees for both "
        "(a) doing their role and (b) achieving results."
    )
    p = doc.add_paragraph("Base pay: fixed salary paid regularly (weekly/monthly).")
    p.style = "List Bullet"
    p = doc.add_paragraph("Variable pay: performance-linked earnings such as commission and bonuses.")
    p.style = "List Bullet"

    doc.add_heading("2) Why Companies Use Sales Compensation", level=1)
    p = doc.add_paragraph("Motivation: reward outcomes, not only activity.")
    p.style = "List Bullet"
    p = doc.add_paragraph("Alignment: connect seller behavior to company goals (growth, renewals, margins).")
    p.style = "List Bullet"
    p = doc.add_paragraph("Fairness: define clear rules so payout is predictable and transparent.")
    p.style = "List Bullet"

    doc.add_heading("3) Core Building Blocks", level=1)
    p = doc.add_paragraph("Quota/Target: expected sales result for a time period.")
    p.style = "List Bullet"
    p = doc.add_paragraph("Attainment: actual results divided by quota.")
    p.style = "List Bullet"
    p = doc.add_paragraph("Commission Rate: percentage used to calculate variable pay.")
    p.style = "List Bullet"
    p = doc.add_paragraph("Accelerator: higher commission rate after crossing quota.")
    p.style = "List Bullet"
    p = doc.add_paragraph("Credit Split: shared deal credit when multiple roles contribute.")
    p.style = "List Bullet"
    p = doc.add_paragraph("Payout Formula: base salary + variable pay components.")
    p.style = "List Bullet"

    doc.add_heading("4) How It Functions (Simple Process)", level=1)
    steps = [
        "Step 1: Plan is designed (quota, rates, bonus rules, payout timing).",
        "Step 2: Deals are closed and credited to reps.",
        "Step 3: Attainment is calculated against quota.",
        "Step 4: Commission is calculated using plan rules (flat or tiered).",
        "Step 5: Bonus/accelerator adjustments are applied.",
        "Step 6: Final payout is audited and paid.",
    ]
    for step in steps:
        p = doc.add_paragraph(step)
        p.style = "List Number"

    doc.add_heading("5) Example A: Basic Monthly Plan (Flat Rate)", level=1)
    base_a = 3500.0
    sales_a = 48000.0
    rate_a = 0.05
    bonus_a = 400.0 if sales_a >= 45000 else 0.0
    commission_a = sales_a * rate_a
    total_a = base_a + commission_a + bonus_a

    table_a = doc.add_table(rows=1, cols=2)
    table_a.style = "Light Grid Accent 1"
    table_a.rows[0].cells[0].text = "Item"
    table_a.rows[0].cells[1].text = "Value"

    rows_a = [
        ("Base salary", usd(base_a)),
        ("Sales closed", usd(sales_a)),
        ("Commission rate", pct(rate_a)),
        ("Commission earned", f"{usd(commission_a)} = {usd(sales_a)} x {pct(rate_a)}"),
        ("Bonus rule", f"{usd(bonus_a)} (because sales >= {usd(45000.0)})"),
        ("Total monthly pay", usd(total_a)),
    ]
    for item, value in rows_a:
        r = table_a.add_row().cells
        r[0].text = item
        r[1].text = value

    doc.add_paragraph(
        "Check: Total Pay = Base + Commission + Bonus = "
        f"{usd(base_a)} + {usd(commission_a)} + {usd(bonus_a)} = {usd(total_a)}"
    )

    doc.add_heading("6) Example B: Tiered Quarterly Plan (More Complex)", level=1)
    base_b = 12000.0
    quota_b = 300000.0
    sales_b = 370000.0

    tier1_cap = 0.80 * quota_b
    tier2_cap = quota_b

    tier1_sales = min(sales_b, tier1_cap)
    tier2_sales = max(0.0, min(sales_b, tier2_cap) - tier1_cap)
    tier3_sales = max(0.0, sales_b - tier2_cap)

    tier1_rate = 0.04
    tier2_rate = 0.06
    tier3_rate = 0.09

    tier1_comm = tier1_sales * tier1_rate
    tier2_comm = tier2_sales * tier2_rate
    tier3_comm = tier3_sales * tier3_rate

    commission_b = tier1_comm + tier2_comm + tier3_comm
    attainment_b = sales_b / quota_b
    bonus_b = 2000.0 if attainment_b >= 1.20 else 0.0
    total_b = base_b + commission_b + bonus_b

    doc.add_paragraph(
        "Plan rules: 4% for 0-80% of quota, 6% for 80-100%, and 9% above 100%. "
        "Bonus = $2,000 if attainment is at least 120%."
    )

    table_b = doc.add_table(rows=1, cols=4)
    table_b.style = "Light Grid Accent 1"
    table_b.rows[0].cells[0].text = "Tier"
    table_b.rows[0].cells[1].text = "Sales in Tier"
    table_b.rows[0].cells[2].text = "Rate"
    table_b.rows[0].cells[3].text = "Commission"

    tier_rows = [
        ("0%-80% of quota", tier1_sales, tier1_rate, tier1_comm),
        ("80%-100% of quota", tier2_sales, tier2_rate, tier2_comm),
        ("Above 100%", tier3_sales, tier3_rate, tier3_comm),
    ]
    for label, sales_in_tier, rate, comm in tier_rows:
        r = table_b.add_row().cells
        r[0].text = label
        r[1].text = usd(sales_in_tier)
        r[2].text = pct(rate)
        r[3].text = usd(comm)

    summary_b = doc.add_table(rows=1, cols=2)
    summary_b.style = "Light Grid Accent 1"
    summary_b.rows[0].cells[0].text = "Metric"
    summary_b.rows[0].cells[1].text = "Value"

    summary_rows = [
        ("Quarterly quota", usd(quota_b)),
        ("Quarterly sales", usd(sales_b)),
        ("Attainment", pct(attainment_b)),
        ("Total commission", usd(commission_b)),
        ("Performance bonus", usd(bonus_b)),
        ("Quarterly base salary", usd(base_b)),
        ("Final quarterly pay", usd(total_b)),
    ]
    for metric, value in summary_rows:
        r = summary_b.add_row().cells
        r[0].text = metric
        r[1].text = value

    doc.add_paragraph(
        "Check: Final Pay = Base + (Tier1 + Tier2 + Tier3 Commission) + Bonus = "
        f"{usd(base_b)} + ({usd(tier1_comm)} + {usd(tier2_comm)} + {usd(tier3_comm)}) + {usd(bonus_b)} = {usd(total_b)}"
    )

    doc.add_heading("7) Example C: Split Credit on One Deal", level=1)
    deal_c = 100000.0
    ae_split = 0.70
    se_split = 0.30
    ae_rate = 0.08
    se_rate = 0.02

    ae_credit = deal_c * ae_split
    se_credit = deal_c * se_split
    ae_payout = ae_credit * ae_rate
    se_payout = se_credit * se_rate

    table_c = doc.add_table(rows=1, cols=5)
    table_c.style = "Light Grid Accent 1"
    table_c.rows[0].cells[0].text = "Role"
    table_c.rows[0].cells[1].text = "Deal Credit %"
    table_c.rows[0].cells[2].text = "Credited Revenue"
    table_c.rows[0].cells[3].text = "Commission Rate"
    table_c.rows[0].cells[4].text = "Variable Pay"

    c_rows = [
        ("Account Executive", ae_split, ae_credit, ae_rate, ae_payout),
        ("Sales Engineer", se_split, se_credit, se_rate, se_payout),
    ]
    for role, split, credited, rate, payout in c_rows:
        r = table_c.add_row().cells
        r[0].text = role
        r[1].text = pct(split)
        r[2].text = usd(credited)
        r[3].text = pct(rate)
        r[4].text = usd(payout)

    doc.add_paragraph(
        "This shows how one deal can be shared fairly across roles. "
        "It prevents overpaying one role and underpaying another."
    )

    doc.add_heading("8) Common Mistakes to Avoid", level=1)
    mistakes = [
        "Using unclear definitions for quota and attainment.",
        "Changing rates mid-period without clear communication.",
        "Ignoring split-credit rules for team selling.",
        "Not auditing payouts before payroll.",
    ]
    for item in mistakes:
        p = doc.add_paragraph(item)
        p.style = "List Bullet"

    doc.add_heading("9) Quick Recap", level=1)
    doc.add_paragraph(
        "Sales compensation works when rules are simple, measurable, and auditable. "
        "The foundation is always the same: clear targets, clear rates, accurate credits, "
        "and correct payout math."
    )

    doc.save(output_path)


if __name__ == "__main__":
    build_document("docs/Sales_Compensation_Basics.docx")
