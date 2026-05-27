from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches


DOC_PATH = Path("docs/Sales_Compensation_Basics.docx")
IMG_PATH = Path("docs/low_level_project_design.png")


def _load_font(size: int):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Helvetica.ttf",
        "/Library/Fonts/Arial.ttf",
    ]
    for path in candidates:
        try:
            return ImageFont.truetype(path, size=size)
        except Exception:
            continue
    return ImageFont.load_default()


def _draw_box(draw, bbox, title, lines, fill, outline, title_font, body_font):
    draw.rounded_rectangle(bbox, radius=16, fill=fill, outline=outline, width=3)
    x1, y1, x2, y2 = bbox

    draw.text((x1 + 16, y1 + 12), title, font=title_font, fill=(20, 20, 20))

    y = y1 + 54
    for line in lines:
        draw.text((x1 + 16, y), f"- {line}", font=body_font, fill=(35, 35, 35))
        y += 34


def _arrow(draw, start, end, color=(70, 70, 70), width=4):
    sx, sy = start
    ex, ey = end
    draw.line([sx, sy, ex, ey], fill=color, width=width)

    head = 12
    if abs(ex - sx) >= abs(ey - sy):
        if ex >= sx:
            points = [(ex, ey), (ex - head, ey - head // 2), (ex - head, ey + head // 2)]
        else:
            points = [(ex, ey), (ex + head, ey - head // 2), (ex + head, ey + head // 2)]
    else:
        if ey >= sy:
            points = [(ex, ey), (ex - head // 2, ey - head), (ex + head // 2, ey - head)]
        else:
            points = [(ex, ey), (ex - head // 2, ey + head), (ex + head // 2, ey + head)]
    draw.polygon(points, fill=color)


def _label(draw, xy, text, font):
    draw.text(xy, text, font=font, fill=(85, 85, 85))


def build_low_level_image(output_path: Path) -> None:
    width, height = 2600, 1750
    img = Image.new("RGB", (width, height), (248, 250, 252))
    draw = ImageDraw.Draw(img)

    title_font = _load_font(48)
    section_font = _load_font(28)
    body_font = _load_font(24)
    small_font = _load_font(22)

    draw.text((640, 26), "Sales Analytics AI - Low-Level Design", font=title_font, fill=(24, 36, 60))

    frontend_box = (110, 130, 860, 430)
    api_box = (930, 130, 1690, 430)
    analytics_box = (110, 540, 860, 940)
    forecasting_box = (930, 540, 1690, 940)
    ingestion_box = (1750, 540, 2490, 940)
    persistence_box = (470, 1080, 2030, 1410)
    quality_box = (1750, 130, 2490, 430)

    _draw_box(
        draw,
        frontend_box,
        "Frontend Components",
        [
            "frontend/src/App.jsx (tab orchestration)",
            "frontend/src/pages/MLInsightsPage.jsx",
            "frontend/src/hooks/useFetch.js",
            "Frontend sends control parameters",
            "Receives KPI, forecast, and warning payloads",
        ],
        fill=(226, 240, 255),
        outline=(56, 119, 194),
        title_font=section_font,
        body_font=body_font,
    )

    _draw_box(
        draw,
        api_box,
        "FastAPI Router Layer",
        [
            "backend/routers/analytics.py",
            "backend/routers/forecasting.py",
            "backend/routers/ingestion.py",
            "backend/routers/agent.py",
            "Dependency injection via get_db AsyncSession",
        ],
        fill=(232, 252, 239),
        outline=(34, 139, 94),
        title_font=section_font,
        body_font=body_font,
    )

    _draw_box(
        draw,
        analytics_box,
        "Analytics and Metric Engine",
        [
            "backend/metrics/calculators.py",
            "Period semantics for closed/open metrics",
            "Rep/team/region filters normalization",
            "Shared metric outputs for API consistency",
            "Warnings surfaced for missing/invalid scope",
        ],
        fill=(255, 244, 229),
        outline=(191, 120, 38),
        title_font=section_font,
        body_font=body_font,
    )

    _draw_box(
        draw,
        forecasting_box,
        "Forecasting and ML Services",
        [
            "backend/ml/forecasting_lab/service.py",
            "models.py candidate model adapters",
            "model_selection.py adaptive ranking",
            "Scenario generation: base/optimistic/conservative",
            "Backtest metrics and model leaderboard",
        ],
        fill=(245, 237, 255),
        outline=(119, 83, 181),
        title_font=section_font,
        body_font=body_font,
    )

    _draw_box(
        draw,
        ingestion_box,
        "Ingestion and Data Build",
        [
            "backend/ingestion/intelligent_ingestion.py",
            "manifest_loader.py canonical mapping",
            "backend/data_generator.py dataset generation",
            "Deterministic payout reconcile + validation",
            "Load modes: full_reload / upsert / append",
        ],
        fill=(255, 237, 237),
        outline=(185, 66, 66),
        title_font=section_font,
        body_font=body_font,
    )

    _draw_box(
        draw,
        persistence_box,
        "Persistence and Data Contracts",
        [
            "backend/database.py Async engine/session",
            "backend/models.py ORM entities and constraints",
            "Core tables: deals, revenue, quotas, payouts",
            "ML tables: ml_predictions, model_runs",
            "Consistency source for API + reports + agent",
        ],
        fill=(236, 239, 244),
        outline=(75, 85, 99),
        title_font=section_font,
        body_font=body_font,
    )

    _draw_box(
        draw,
        quality_box,
        "Quality and Audit Controls",
        [
            "backend/validation/revops_rules.py",
            "backend/audit/payout_audit.py",
            "Cent-level payout tolerance checks",
            "Fail-fast on missing/mismatched duplicates",
            "Evidence for trust and reproducibility",
        ],
        fill=(255, 247, 221),
        outline=(170, 124, 33),
        title_font=section_font,
        body_font=body_font,
    )

    _arrow(draw, (860, 280), (930, 280))
    _label(draw, (870, 238), "HTTP JSON", small_font)

    _arrow(draw, (1310, 430), (1310, 540))
    _label(draw, (1330, 474), "dispatch", small_font)

    _arrow(draw, (1110, 430), (510, 540))
    _arrow(draw, (1510, 430), (2070, 540))

    _arrow(draw, (530, 940), (900, 1080))
    _arrow(draw, (1310, 940), (1260, 1080))
    _arrow(draw, (2070, 940), (1620, 1080))

    _arrow(draw, (1750, 280), (1690, 280))
    _label(draw, (1710, 238), "checks", small_font)

    draw.text(
        (210, 1480),
        "Main runtime flow: Frontend -> Routers -> Service modules -> ORM/DB -> Response with diagnostics and warnings",
        font=small_font,
        fill=(45, 45, 45),
    )
    draw.text(
        (210, 1520),
        "Data trust flow: Ingestion/Generator -> Validation/Audit -> Deterministic payout math -> Consistent analytics and forecasts",
        font=small_font,
        fill=(45, 45, 45),
    )

    img.save(output_path)


def embed_in_doc(doc_path: Path, image_path: Path) -> None:
    doc = Document(str(doc_path))
    heading = "Appendix: Low-Level Project Design"

    already_exists = any(p.text.strip() == heading for p in doc.paragraphs)
    if not already_exists:
        doc.add_page_break()
        doc.add_heading(heading, level=1)
        doc.add_paragraph(
            "This low-level diagram maps key modules and file-level responsibilities "
            "across frontend, API routers, service logic, and persistence."
        )
        doc.add_picture(str(image_path), width=Inches(6.5))
        doc.add_paragraph("Figure: Low-level architecture and data-flow responsibilities")

    doc.save(str(doc_path))


def main() -> None:
    build_low_level_image(IMG_PATH)
    embed_in_doc(DOC_PATH, IMG_PATH)


if __name__ == "__main__":
    main()
