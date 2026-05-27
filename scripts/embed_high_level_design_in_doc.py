from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches


DOC_PATH = Path("docs/Sales_Compensation_Basics.docx")
IMG_PATH = Path("docs/high_level_project_design.png")


def _load_font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Helvetica.ttf",
        "/Library/Fonts/Arial.ttf",
    ]
    for path in candidates:
        try:
            return ImageFont.truetype(path, size=size)
        except Exception:
            continue
    return ImageFont.load_default()


def _draw_rounded_box(draw: ImageDraw.ImageDraw, xy, fill, outline, radius=18, width=3):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def _centered_text(draw: ImageDraw.ImageDraw, bbox, text: str, font, fill=(20, 20, 20)):
    x1, y1, x2, y2 = bbox
    tw, th = draw.multiline_textbbox((0, 0), text, font=font, spacing=6)[2:]
    cx = x1 + (x2 - x1 - tw) / 2
    cy = y1 + (y2 - y1 - th) / 2
    draw.multiline_text((cx, cy), text, font=font, fill=fill, align="center", spacing=6)


def _arrow(draw: ImageDraw.ImageDraw, start, end, fill=(70, 70, 70), width=5):
    sx, sy = start
    ex, ey = end
    draw.line([sx, sy, ex, ey], fill=fill, width=width)
    # Arrow head
    ah = 14
    if abs(ex - sx) > abs(ey - sy):
        # Horizontal arrow
        if ex >= sx:
            points = [(ex, ey), (ex - ah, ey - ah // 2), (ex - ah, ey + ah // 2)]
        else:
            points = [(ex, ey), (ex + ah, ey - ah // 2), (ex + ah, ey + ah // 2)]
    else:
        # Vertical arrow
        if ey >= sy:
            points = [(ex, ey), (ex - ah // 2, ey - ah), (ex + ah // 2, ey - ah)]
        else:
            points = [(ex, ey), (ex - ah // 2, ey + ah), (ex + ah // 2, ey + ah)]
    draw.polygon(points, fill=fill)


def _label(draw: ImageDraw.ImageDraw, at, text: str, font, fill=(80, 80, 80)):
    draw.text(at, text, font=font, fill=fill)


def build_architecture_image(output_path: Path) -> None:
    width, height = 2200, 1350
    image = Image.new("RGB", (width, height), (248, 250, 252))
    draw = ImageDraw.Draw(image)

    title_font = _load_font(52)
    box_font = _load_font(30)
    label_font = _load_font(24)

    draw.text((520, 26), "Sales Analytics AI - High-Level Architecture", font=title_font, fill=(24, 36, 60))

    # Boxes: (x1, y1, x2, y2)
    frontend = (770, 120, 1450, 255)
    backend = (700, 350, 1520, 540)
    data_services = (170, 650, 930, 840)
    ml_services = (1270, 650, 2030, 840)
    db_layer = (700, 960, 1520, 1145)
    ingestion = (70, 960, 580, 1145)
    llm = (1620, 350, 2110, 540)

    _draw_rounded_box(draw, frontend, fill=(226, 240, 255), outline=(56, 119, 194))
    _draw_rounded_box(draw, backend, fill=(232, 252, 239), outline=(34, 139, 94))
    _draw_rounded_box(draw, data_services, fill=(255, 244, 229), outline=(191, 120, 38))
    _draw_rounded_box(draw, ml_services, fill=(245, 237, 255), outline=(119, 83, 181))
    _draw_rounded_box(draw, db_layer, fill=(236, 239, 244), outline=(75, 85, 99))
    _draw_rounded_box(draw, ingestion, fill=(255, 237, 237), outline=(185, 66, 66))
    _draw_rounded_box(draw, llm, fill=(255, 247, 221), outline=(170, 124, 33))

    _centered_text(draw, frontend, "Frontend Dashboard\nReact + Vite\nCharts + Controls", box_font)
    _centered_text(
        draw,
        backend,
        "Backend API Layer (FastAPI)\nRouters: /analytics, /ml, /agent, /ingestion\nShared business rules + orchestration",
        box_font,
    )
    _centered_text(
        draw,
        data_services,
        "Data & Metrics Services\nCanonical model\nMetric calculators\nValidation + consistency checks",
        box_font,
    )
    _centered_text(
        draw,
        ml_services,
        "ML Services\nForecasting Lab\nDeal Scoring\nRep Clustering\nBacktest + model selection",
        box_font,
    )
    _centered_text(
        draw,
        db_layer,
        "PostgreSQL + SQLAlchemy\nDeals, revenue, quotas, payouts, features\nSingle source of truth",
        box_font,
    )
    _centered_text(
        draw,
        ingestion,
        "Ingestion + Quality Pipeline\nCSV/Manifest mapping\nLoad modes: reload/upsert/append\nData quality checks",
        box_font,
    )
    _centered_text(draw, llm, "Optional LLM Provider\nUsed by agent workflows", box_font)

    # Arrows
    _arrow(draw, (1110, 255), (1110, 350))
    _label(draw, (1130, 290), "HTTP/JSON", label_font)

    _arrow(draw, (910, 540), (620, 650))
    _arrow(draw, (1310, 540), (1580, 650))

    _arrow(draw, (550, 960), (700, 1060))
    _label(draw, (520, 920), "validated load", label_font)

    _arrow(draw, (550, 1060), (700, 1060))

    _arrow(draw, (560, 840), (930, 960))
    _arrow(draw, (1650, 840), (1285, 960))
    _label(draw, (860, 900), "read/write", label_font)
    _label(draw, (1450, 900), "train/infer", label_font)

    _arrow(draw, (1520, 430), (1620, 430))
    _label(draw, (1536, 390), "agent calls", label_font)

    image.save(output_path)


def embed_image_in_doc(doc_path: Path, image_path: Path) -> None:
    doc = Document(str(doc_path))

    heading_exists = any(p.text.strip() == "Appendix: High-Level Project Design" for p in doc.paragraphs)
    if not heading_exists:
        doc.add_page_break()
        doc.add_heading("Appendix: High-Level Project Design", level=1)
        doc.add_paragraph(
            "This diagram summarizes the major layers of the Sales Analytics AI platform "
            "and how data and requests flow across them."
        )
        doc.add_picture(str(image_path), width=Inches(6.5))
        doc.add_paragraph("Figure: High-level architecture of the project")

    doc.save(str(doc_path))


def main() -> None:
    build_architecture_image(IMG_PATH)
    embed_image_in_doc(DOC_PATH, IMG_PATH)


if __name__ == "__main__":
    main()
